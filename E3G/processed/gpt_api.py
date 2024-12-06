import sys
from openai import AzureOpenAI
import os
from dotenv import load_dotenv
from docx import Document
import pdfplumber
import csv
import json
import openpyxl

def gpt4(system, user):
    # Load environment file for secrets.
    try:
        if load_dotenv('.env') is False:
            raise TypeError
    except TypeError:
        print('Unable to load .env file.')
        quit()
    api_key = os.getenv('OPENAI_API_KEY')
    api_base = os.getenv('OPENAI_API_BASE')
    organization = os.getenv('OPENAI_ORGANIZATION')
    api_version = os.getenv('API_VERSION')
    model = os.getenv('MODEL')
    print(api_key)
    print(api_base)
    print(organization)
    print(api_version)
    print(model)
    # Create Azure client
    client = AzureOpenAI(
        api_key=api_key,  
        api_version=api_version,
        azure_endpoint=api_base,
        organization=organization
    )

    # Send a completion call to generate an answer
    print('Sending a test completion job')

    response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": f"{system}"},
                {"role": "user", "content": f"{user}"}
            ],
            temperature=0,
            stop=None)
    print("API call success!")
    response = response.choices[0].message.content
    return response

def embedding(input, dimension):
    # Load environment file for secrets.
    try:
        if load_dotenv('.env') is False:
            raise TypeError
    except TypeError:
        print('Unable to load .env file.')
        quit()
    api_key = os.getenv('OPENAI_API_KEY')
    api_base = os.getenv('OPENAI_API_BASE')
    organization = os.getenv('OPENAI_ORGANIZATION')
    api_version = os.getenv('API_VERSION')
    embedding_model = os.getenv('EMBEDDING_MODEL')
    # Create Azure client
    client = AzureOpenAI(
        api_key=api_key,  
        api_version=api_version,
        azure_endpoint=api_base,
        organization=organization
    )

    # Send a completion call to generate an answer
    print('Sending a test completion job')

    response = client.embeddings.create(
        input=input,
        model=embedding_model,
    )
    return response

def read_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
    return text

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables.append(table_data)
    return '\n'.join(full_text), tables

def read_csv(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as csvfile:
            csv_read = csv.reader(csvfile)
            first_line = next(csv_read)
            return first_line
    except Exception as e:
        print(f"Error reading CSV file {file_path}: {e}")
    return None

def read_excel(file_path):
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        rows = list(sheet.iter_rows(values_only=True))
        result = []
        for idx, row in enumerate(rows):
            row_text = f"{'first' if idx == 0 else 'second' if idx == 1 else f'{idx+1}th'} row: {', '.join(map(str, row))}"
            result.append(row_text)
        return "; ".join(result)
    except Exception as e:
        print(f"Error reading Excel file {file_path}: {e}")
        return None

def write_json(data, output_file):
    try:
        with open(output_file, 'w', encoding='utf-8') as json_file:
            json.dump(data, json_file, ensure_ascii=False, indent=4)
    except Exception as e:
        print(f"Error writing to JSON file {output_file}: {e}")

def main(des_path, csv_path, json_output_path):
    # Read content from files
    col = read_csv(csv_path)

    # Determine the description file type and read its content
    if des_path.endswith('.pdf'):
        des = read_pdf(des_path)
    elif des_path.endswith('.docx'):
        print("reading doc")
        des = read_docx(des_path)
    elif des_path.endswith('.xlsx'):
        des = read_excel(des_path)
    else:
        des = None

    if des is not None:
        # Get the response from GPT-4
        result = gpt4("Combine everything together, to make a description for this dataset in three or four sentenses", f"col: {col}; des: {des}")

        # Write the result to a JSON file
        write_json(result, json_output_path)

        print(result)
    else:
        print("Unsupported file type for description file.")

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python script.py <des_path> <csv_path> <json_output_path>")
        sys.exit(1)
    
    des_path = sys.argv[1]
    csv_path = sys.argv[2]
    json_output_path = sys.argv[3]
    
    main(des_path, csv_path, json_output_path)
